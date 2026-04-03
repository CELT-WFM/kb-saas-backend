Here is your FastAPI router code formatted correctly for readability. I've organized the imports, added proper spacing between functions, and cleaned up the indentation so it’s easy to debug.

Python

import os
import boto3
import openai
import pandas as pd
from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from sqlalchemy.orm import Session
from pypdf import PdfReader
from docx import Document as DocxDocument

# Local imports
from db import get_db
from models import Document, Chunk
from auth import verify_token

router = APIRouter()

# AWS S3 Configuration
s3 = boto3.client(
    "s3",
    aws_access_key_id=os.getenv("S3_ACCESS_KEY"),
    aws_secret_access_key=os.getenv("S3_SECRET_KEY"),
    endpoint_url=os.getenv("S3_ENDPOINT")
)
BUCKET = os.getenv("S3_BUCKET")

# OpenAI Configuration
client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def chunk_text(text, size=800):
    words = text.split()
    return [" ".join(words[i:i+size]) for i in range(0, len(words), size)]

def embed(text):
    r = client.embeddings.create(model="text-embedding-3-small", input=text)
    # Note: Ensure your DB 'embedding' column type matches this return (string vs vector)
    return str(r.data[0].embedding)

def extract_text(file):
    ext = file.filename.split(".")[-1].lower()
    
    if ext == "pdf":
        reader = PdfReader(file.file)
        return " ".join(p.extract_text() or "" for p in reader.pages)
    
    elif ext in ["xls", "xlsx"]:
        return pd.read_excel(file.file).to_string()
    
    elif ext == "csv":
        return pd.read_csv(file.file).to_string()
    
    elif ext in ["docx", "doc"]:
        doc = DocxDocument(file.file)
        return " ".join(p.text for p in doc.paragraphs)
    
    raise HTTPException(status_code=400, detail="Unsupported file type")

@router.post("/upload")
async def upload(
    collection_id: int, 
    file: UploadFile = File(...),
    user_id: int = Depends(verify_token), 
    db: Session = Depends(get_db)
):
    # 1. Upload to S3
    s3_key = f"{user_id}/{file.filename}"
    s3.upload_fileobj(file.file, BUCKET, s3_key)
    
    # 2. Reset file pointer and extract text
    file.file.seek(0)
    text = extract_text(file)
    
    # 3. Create Document Record
    doc = Document(
        filename=file.filename, 
        s3_key=s3_key,
        owner_id=user_id, 
        collection_id=collection_id,
        file_type=file.filename.split(".")[-1].lower()
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    # 4. Chunk and Embed
    for chunk in chunk_text(text):
        c = Chunk(
            document_id=doc.id, 
            collection_id=collection_id,
            content=chunk, 
            embedding=embed(chunk)
        )
        db.add(c)
    
    db.commit()
    return {"document_id": doc.id, "status": "indexed"}
